"""
Classwork Studio - 香港教育课堂作业生成器
Flask 后端主文件，集成 MiniMax API（兼容 OpenAI 接口）
"""

import os
import io
import sys
import re
import time
from flask import Flask, request, jsonify, send_from_directory, send_file, Response, session, redirect, url_for
from flask_cors import CORS
from dotenv import load_dotenv
from openai import OpenAI
import PyPDF2

# 修复旧版 httpx 的 proxies 参数兼容性问题
# 新版 httpx 0.28+ 使用 proxy（单数），旧版使用 proxies（复数）
try:
    import httpx
    _orig_client_init = httpx.Client.__init__

    def _patched_init(self, *args, **kwargs):
        # 如果传入了 proxies 但不是 proxy，转换它
        if 'proxies' in kwargs and 'proxy' not in kwargs:
            kwargs['proxy'] = kwargs.pop('proxies')
        # 如果同时传入了 proxy 和 proxies，移除 proxies
        if 'proxies' in kwargs and 'proxy' in kwargs:
            kwargs.pop('proxies')
        return _orig_client_init(self, *args, **kwargs)

    httpx.Client.__init__ = _patched_init
except Exception:
    pass

# 加载 .env 文件中的环境变量
load_dotenv()

# 初始化 Flask 应用，静态资源目录指向 static 文件夹
app = Flask(__name__, static_folder="static", static_url_path="/static")
CORS(app, origins=["http://localhost:3000", "http://127.0.0.1:3000"])

# Session 配置（用于登录认证）
app.secret_key = os.getenv("SECRET_KEY", "dev_secret_key")
app.config["SESSION_TYPE"] = "filesystem"
app.config["SESSION_COOKIE_SECURE"] = False
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"

# 从 .env 加载（通用变量，不再硬编码 provider 名称）
# 任意兼容 OpenAI Chat Completions 接口的模型服务都可以用。
# 在 .env 里配置一个或多个 LLM_API_KEY_<MODEL> 即可。
# 详见 .env.example 与 README。

def _load_model_config():
    """
    Build MODEL_CONFIG dynamically from environment variables.

    Recognised variables for each entry:
        LLM_<KEY>_API_KEY      required
        LLM_<KEY>_BASE_URL     optional, default https://api.openai.com/v1
        LLM_<KEY>_MODEL        optional, default = <KEY>
        LLM_<KEY>_MAX_TOKENS   optional, default 32768

    Where <KEY> is the human-friendly identifier used in the front-end
    selector, e.g.  LLM_PROD_API_KEY=...   ->  selector key "PROD".

    Fallback order (for backwards compatibility with earlier checkouts):
        LLM_API_KEY  -> single-model mode; key = "default"
    """
    import re as _re_key
    cfg = {}

    # ---- Multi-model mode: scan LLM_<KEY>_* variables ----
    env = os.environ
    pattern = _re_key.compile(r"^LLM_([A-Z0-9_]+)_API_KEY$")
    keys_by_name = {}
    for env_key, value in env.items():
        m = pattern.match(env_key)
        if not m or not value.strip():
            continue
        name = m.group(1)
        keys_by_name[name] = value.strip()

    for name in sorted(keys_by_name):
        prefix = f"LLM_{name}_"
        cfg[name] = {
            "api_key":     keys_by_name[name],
            "base_url":    env.get(prefix + "BASE_URL",   "https://api.openai.com/v1").strip(),
            "model_name":  env.get(prefix + "MODEL",      name).strip(),
            "max_tokens":  int(env.get(prefix + "MAX_TOKENS", "32768").strip() or 32768),
        }

    # ---- Single-model fallback ----
    if not cfg:
        legacy_key = env.get("LLM_API_KEY", "").strip()
        if legacy_key:
            cfg["default"] = {
                "api_key":     legacy_key,
                "base_url":    env.get("LLM_BASE_URL",   "https://api.openai.com/v1").strip(),
                "model_name":  env.get("LLM_MODEL",      "default").strip(),
                "max_tokens":  int(env.get("LLM_MAX_TOKENS", "32768").strip() or 32768),
            }

    if not cfg:
        raise RuntimeError(
            "未配置任何模型 API Key。\n"
            "请在项目根目录创建 .env 文件，并按以下任一方式配置：\n"
            "  - 单模型：  LLM_API_KEY=sk-...\n"
            "  - 多模型：  LLM_<NAME>_API_KEY=sk-...  (NAME 是任意字母数字下划线)\n"
            "            例如 LLM_PROD_API_KEY / LLM_FAST_API_KEY\n"
            "详见 .env.example。"
        )

    return cfg

MODEL_CONFIG = _load_model_config()

# Default model = MODEL_NAME env var, or the first key in MODEL_CONFIG.
DEFAULT_MODEL = os.getenv("MODEL_NAME") or next(iter(MODEL_CONFIG.keys()))

# 管理员账号配置
ADMIN_USERNAME = "admin"
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "teacher@2026")

def get_model_client(model_name: str = None):
    """Get OpenAI-compatible client for the specified model."""
    model_name = model_name or DEFAULT_MODEL
    config = MODEL_CONFIG.get(model_name)
    if not config:
        raise ValueError(f"Unknown model: {model_name}")
    if not config["api_key"]:
        raise ValueError(f"API key not configured for model: {model_name}")
    return OpenAI(
        api_key=config["api_key"],
        base_url=config["base_url"]
    )

def get_model_max_tokens(model_name: str = None) -> int:
    """Get max_tokens for the specified model."""
    model_name = model_name or DEFAULT_MODEL
    config = MODEL_CONFIG.get(model_name)
    return config["max_tokens"] if config else 32768

def get_model_name(model_name: str = None) -> str:
    """Get the actual API model name."""
    model_name = model_name or DEFAULT_MODEL
    config = MODEL_CONFIG.get(model_name)
    return config["model_name"] if config else model_name

# 初始化默认模型客户端（兼容任意 OpenAI Chat Completions 接口的服务）
# 注意：实际生成路径走 get_model_client()，这个 client 仅供兜底；保留它
# 是为了兼容任何直接 import `client` 的代码路径。
_default_cfg = MODEL_CONFIG[DEFAULT_MODEL]
client = OpenAI(
    api_key=_default_cfg["api_key"],
    base_url=_default_cfg["base_url"],
)

# Curriculum folder path
CURRICULUM_FOLDER = os.path.join(os.path.dirname(__file__), "curriculum")

# Task storage folder (for progress persistence)
TASKS_FOLDER = os.path.join(os.path.dirname(__file__), "tasks")
os.makedirs(TASKS_FOLDER, exist_ok=True)

# Level name to code mapping
LEVEL_CODES = {
    "Primary 1": "P1", "Primary 2": "P2", "Primary 3": "P3",
    "Primary 4": "P4", "Primary 5": "P5", "Primary 6": "P6",
    "Secondary 1": "S1", "Secondary 2": "S2", "Secondary 3": "S3",
    "Secondary 4": "S4", "Secondary 5": "S5", "Secondary 6": "S6",
}

# Subject name to code mapping
SUBJECT_CODES = {
    "Chinese Language": "Chinese",
    "English Language": "English",
    "Mathematics": "Maths",
    "General Studies": "GST",
    "Putonghua": "Putonghua",
    "Music": "Music",
    "Physical Education": "PE",
    "Visual Arts": "VA",
    "Computer Studies": "Computer",
    "Physics": "Physics",
    "Chemistry": "Chemistry",
    "Biology": "Biology",
    "History": "History",
    "Geography": "Geography",
    "Economics": "Economics",
    "BAFS (Accounting)": "BAFS_Acc",
    "BAFS (Business)": "BAFS_Bus",
    "Information and Communication Technology": "ICT",
    "Citizenship and Social Development": "CSD",
}


def get_curriculum_context(filename):
    """
    Read a specific curriculum PDF file by filename.
    Returns the full extracted text (no character limit).
    """
    if not filename:
        return ""

    if not os.path.exists(CURRICULUM_FOLDER):
        return ""

    pdf_path = os.path.join(CURRICULUM_FOLDER, filename)

    if not os.path.exists(pdf_path):
        return ""

    try:
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            full_text = "\n".join(text_parts)
            context = f"[{filename}]\n{full_text}"
            return f"\n\n=== CURRICULUM REFERENCE ===\n{context}\n=== END CURRICULUM REFERENCE ===\n"
    except Exception as e:
        print(f"[WARN] Failed to read curriculum PDF {filename}: {e}")
        return ""


def get_available_curriculum_files():
    """Return list of available curriculum PDF files."""
    if not os.path.exists(CURRICULUM_FOLDER):
        return []
    files = [f for f in os.listdir(CURRICULUM_FOLDER) if f.endswith(".pdf")]
    return sorted(files)


import json as _json
import uuid as _uuid

def _get_task_path(task_id: str) -> str:
    return os.path.join(TASKS_FOLDER, f"{task_id}.json")

def _save_task(task_id: str, data: dict, user: str = None):
    path = _get_task_path(task_id)
    data["updated_at"] = time.time()
    if user:
        data["user"] = user
    try:
        with open(path, "w", encoding="utf-8") as f:
            _json.dump(data, f, ensure_ascii=False)
        print(f"[DEBUG] Task saved: {path}, user={user}, status={data.get('status')}")
    except Exception as e:
        print(f"[ERROR] Failed to save task {task_id}: {e}, path={path}")

def _load_task(task_id: str) -> dict | None:
    path = _get_task_path(task_id)
    if not os.path.exists(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return _json.load(f)
    except Exception:
        return None

def _delete_task(task_id: str):
    path = _get_task_path(task_id)
    try:
        os.remove(path)
    except Exception:
        pass

def _cleanup_expired_tasks(max_age_hours: int = 24):
    """清理超过指定小时数的过期任务。"""
    if not os.path.exists(TASKS_FOLDER):
        return 0
    now = time.time()
    expired_count = 0
    max_age_seconds = max_age_hours * 3600
    for fname in os.listdir(TASKS_FOLDER):
        if not fname.endswith(".json"):
            continue
        task_id = fname[:-5]
        path = _get_task_path(task_id)
        try:
            with open(path, "r", encoding="utf-8") as f:
                task = _json.load(f)
                updated_at = task.get("updated_at", 0)
                if now - updated_at > max_age_seconds:
                    os.remove(path)
                    expired_count += 1
        except Exception:
            try:
                os.remove(path)
                expired_count += 1
            except Exception:
                pass
    if expired_count > 0:
        print(f"[CLEANUP] Removed {expired_count} expired task(s)")
    return expired_count


def _start_cleanup_thread():
    """启动后台线程，定期清理过期任务（每小时执行一次）。"""
    import threading
    
    def cleanup_loop():
        while True:
            _cleanup_expired_tasks(24)
            time.sleep(3600)
    
    thread = threading.Thread(target=cleanup_loop, daemon=True)
    thread.start()
    print("[CLEANUP] Background cleanup thread started")


def _list_tasks(user: str = None) -> list:
    if not os.path.exists(TASKS_FOLDER):
        return []
    tasks = []
    for fname in os.listdir(TASKS_FOLDER):
        if not fname.endswith(".json"):
            continue
        task_id = fname[:-5]
        task = _load_task(task_id)
        if task:
            task_user = task.get("user")
            if user and task_user and task_user != user:
                continue
            tasks.append({
                "task_id": task_id,
                "status": task.get("status", "unknown"),
                "level": task.get("level", ""),
                "subject": task.get("subject", ""),
                "progress": task.get("progress", 0),
                "char_count": len(task.get("content", "")),
                "created_at": task.get("created_at", 0),
                "updated_at": task.get("updated_at", 0),
            })
    return sorted(tasks, key=lambda t: t["updated_at"], reverse=True)


def requires_login(f):
    """Decorator to protect routes that require authentication."""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user" not in session:
            return jsonify({"error": "Unauthorized", "message": "Please login first"}), 401
        return f(*args, **kwargs)
    return decorated_function


@app.route("/login")
def login_page():
    if "user" in session:
        return redirect("/")
    return send_from_directory("static", "login.html")


@app.route("/api/login", methods=["POST"])
def api_login():
    try:
        data = request.get_json(silent=True) or {}
        username = (data.get("username") or "").strip()
        password = (data.get("password") or "").strip()

        if not username or not password:
            return jsonify({"success": False, "message": "Username and password are required"}), 400

        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            session["user"] = {"username": username}
            return jsonify({"success": True, "message": "Login successful"}), 200

        return jsonify({"success": False, "message": "Invalid username or password"}), 401
    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500


@app.route("/api/logout", methods=["POST"])
def api_logout():
    session.clear()
    return jsonify({"success": True, "message": "Logged out successfully"}), 200


@app.route("/api/check-login", methods=["GET"])
def check_login():
    if "user" in session:
        return jsonify({"loggedIn": True, "user": session["user"]}), 200
    return jsonify({"loggedIn": False}), 200


@app.route("/api/models", methods=["GET"])
def get_models():
    models = []
    for name, config in MODEL_CONFIG.items():
        models.append({
            "name": name,
            "max_tokens": config["max_tokens"],
            "configured": bool(config["api_key"]),
        })
    return jsonify({"models": models, "default": DEFAULT_MODEL}), 200


@app.route("/api/tasks", methods=["GET"])
@requires_login
def list_tasks():
    user = session.get("user", {}).get("username", "")
    tasks = _list_tasks(user)
    return jsonify({"tasks": tasks}), 200


@app.route("/api/tasks/<task_id>", methods=["GET"])
@requires_login
def get_task(task_id):
    task = _load_task(task_id)
    if not task:
        return jsonify({"error": "Task not found"}), 404
    return jsonify(task), 200


@app.route("/api/tasks/<task_id>", methods=["DELETE"])
@requires_login
def delete_task(task_id):
    _delete_task(task_id)
    return jsonify({"success": True, "message": "Task deleted"}), 200


@app.route("/api/tasks/<task_id>/resume", methods=["POST"])
@requires_login
def resume_task(task_id):
    task = _load_task(task_id)
    if not task:
        return jsonify({"error": "Task not found"}), 404
    
    status = task.get("status")
    if status == "completed":
        return jsonify({
            "status": "completed",
            "content": task.get("content", ""),
            "tokens": task.get("tokens", {}),
            "time": task.get("elapsed", 0),
            "finish_reason": task.get("finish_reason", "done"),
        }), 200
    
    return jsonify({
        "status": status,
        "content": task.get("content", ""),
        "progress": task.get("progress", 0),
        "elapsed": task.get("elapsed", 0),
    }), 200


# 根路由：返回前端页面
@app.route("/")
def index():
    if "user" not in session:
        return redirect("/login")
    return send_from_directory("static", "index.html")


# 获取可用的课程文件列表
@requires_login
@app.route("/api/curriculum/files", methods=["GET"])
def curriculum_files():
    files = get_available_curriculum_files()
    return jsonify({"files": files, "count": len(files)})


# 课堂作业生成接口
@requires_login
@app.route("/api/generate", methods=["POST"])
def generate_classwork():
    try:
        # 解析请求体
        data = request.get_json(silent=True) or {}
        level = (data.get("level") or "").strip()
        subject = (data.get("subject") or "").strip()
        difficulty = (data.get("difficulty") or "").strip()
        curriculum_file = (data.get("curriculumFile") or "").strip()
        chinese_version = data.get("chineseVersion", False)
        overall_notes = (data.get("overallNotes") or "").strip()
        model_name = (data.get("model") or "").strip() or None

        # 参数校验
        if not level or not subject or not difficulty:
            return jsonify({"error": "缺少必要参数：level、subject、difficulty"}), 400

        # 创建任务
        task_id = str(_uuid.uuid4())
        user = session.get("user", {}).get("username", "")
        print(f"[DEBUG] Simple generate request: level={level}, subject={subject}, user={user}, session_user={session.get('user')}")
        _save_task(task_id, {
            "status": "started",
            "content": "",
            "progress": 0,
            "level": level,
            "subject": subject,
            "created_at": time.time(),
        }, user)

        # 根据年级推断建议时长
        if level.startswith("Primary 1") or level.startswith("Primary 2") or level.startswith("Primary 3"):
            suggested_time = 40
        elif level.startswith("Primary"):
            suggested_time = 50
        elif level.startswith("Secondary 1") or level.startswith("Secondary 2") or level.startswith("Secondary 3"):
            suggested_time = 60
        else:
            suggested_time = 75

        # 获取课程文件上下文（如果有的话）
        curriculum_context = get_curriculum_context(curriculum_file)
        if curriculum_context:
            print(f"[INFO] Loaded curriculum: {curriculum_file}")

        # 系统提示词：香港资深教师/作业设计师，按统一模板生成多部分、带分值的课堂作业
        system_prompt = """You are an experienced Hong Kong primary/secondary school teacher and a professional worksheet designer. Design an original, classroom-ready classwork worksheet with the parameters below.

PARAMETERS
- Education Level: __LEVEL__
- Subject: __SUBJECT__
- Difficulty: __DIFFICULTY__
- Language Mode: __LANGUAGE_MODE__
- Overall Notes: __OVERALL_NOTES__

OUTPUT STRUCTURE (use Markdown, follow this skeleton exactly):

# __LEVEL__ __SUBJECT__ Classwork (Difficulty: __DIFFICULTY__)
**Total Marks: 100** | **Suggested Time: __TIME__ minutes** | **Date: __________** | **Name: __________**

## Learning Objectives
- (3-5 specific, measurable objectives in bullet form, e.g. "Identify and use the past simple tense correctly in context.")

---

## Part I: <Part Title> (X questions * Y marks = Z marks)
<one-line instructions if needed>
1. <question>
2. <question>
...

## Part II: <Part Title> (X questions * Y marks = Z marks)
...

(include 4-7 Parts in total, each clearly named and weighted; Part totals must add up to exactly 100)

---

## Answers
(Provide complete, concise answers in the same order. Format: "Part I, Q1: <answer>". For multi-step answers, show the steps.)

END OF OUTPUT STRUCTURE

DO NOT INCLUDE any sections after "## Answers" (such as Design Rules, instructions, or meta-commentary). The worksheet ends at the Answers section.

---

PRIORITY HIERARCHY (READ FIRST - determines which instructions take precedence):
- **Overall Notes** (highest priority): If the teacher provided overall notes, those instructions apply to the ENTIRE worksheet and override ALL other rules below.
- **System Defaults / DESIGN RULES** (lower priority): Apply only when no conflicting instructions exist in the overall notes.

---

DESIGN RULES (follow these rules internally; do NOT include them in the output):

0. COMPLETENESS (CRITICAL - you MUST generate a FULL, DETAILED worksheet of at least 10,000+ characters total):
   - Each Part MUST contain the FULL set of questions as declared in its heading (e.g. "5 questions * 4 marks = 20 marks" means 5 complete questions, not 1-2).
   - The "## Answers" section MUST contain a numbered answer for EVERY question across ALL Parts. No "..." or placeholders.
   - Reading passages MUST be at least 250-400 words for secondary level, 150-250 for primary.
   - Each short-answer question in the worksheet must have a complete, substantive answer (1-3 sentences) in the answer key.
   - DO NOT summarize, abbreviate, or skip questions. DO NOT use "..." or "etc." in place of full content.
   - Your total output MUST be comprehensive and detailed - aim for 10000-15000+ characters of worksheet content.

1. MARKS: All Part totals must add up to exactly 100. Show "(X qs * Y marks = Z marks)" in each Part heading.

2. DIFFICULTY CALIBRATION (controls cognitive demand):
   - **Basic** - recall and recognition. Short items, single-step, direct application, formulaic, low Bloom's level.
   - **Advanced** - application and analysis. Multi-step, comparison, error correction, sentence transformation, structured data-response.
   - **Expert** - synthesis and evaluation. Extended open response (>=150 words), critical reasoning, original argument, cross-topic integration, real-world application.

3. LEVEL CALIBRATION (controls vocabulary, length, abstraction):
   - **Primary 1-3**: very simple vocabulary, short sentences, concrete everyday topics, 40-50 min content.
   - **Primary 4-6**: age-appropriate vocabulary, multi-clause sentences, 50-60 min content.
   - **Secondary 1-3**: broader vocabulary, abstract concepts, longer passages (250-400 words), structured questions, 60-70 min content.
   - **Secondary 4-6**: academic vocabulary, sophisticated passages, data-response, discursive writing (350+ words), 70-80 min content.

4. SUGGESTED TIME: use the value provided in the header.

5. SUBJECT TEMPLATES (pick 4-7 relevant Parts; each Part = one question type):
   - **English Language**: Reading Comprehension, Vocabulary, Grammar (fill-in-the-blanks), Proofreading/Error Correction, Sentence Transformation, Cloze Test, Writing.
   - **Chinese Language (中国语文)**: 課文理解、字詞解釋(小學)、句式改寫(小學)、標點改錯(小學)、閱讀理解、文言文(中學)、寫作(中學)。Use 繁體中文 for Hong Kong convention.
   - **Mathematics**: Direct Calculation / Multiple Choice / Fill in the Blanks / Word Problems / Show Your Steps / Diagram-based. Use HK conventions ($, cm, kg). Answer key must show step-by-step working.
   - **Sciences (Physics / Chemistry / Biology)**: Concept MCQ / Key-term Fill in the Blanks / Short Structured Questions / Calculation with SI units / Diagram or Experiment description.
   - **Humanities (History / Geography / Economics, BAFS Accounting, BAFS Business)**: MCQ / Data Response / Source-based Questions / Short Structured / Extended Response.
   - **BAFS (Accounting)**: Focus on double-entry bookkeeping, journal entries, ledgers, trial balance, financial statements (Income Statement & Balance Sheet), accounting ratios, depreciation, and interpretation of financial data.
   - **BAFS (Business)**: Focus on business management concepts, entrepreneurship, marketing mix (4Ps), HRM, operations management, business ethics, case studies, and business environment analysis.
   - **Other (Music / PE / Visual Arts / ICT / Computer Studies)**: Knowledge Recall / Practical Application / Mini-project Task.

6. FORMATTING:
   - Use `##` for Parts, `###` for sub-sections, numbered lists for questions.
   - **Multiple Choice**: Format options as `A) B) C) D)` on the SAME line, separated by spaces. 
   - **Fill in the blank (inline)**: Use `__________` (10 underscores) for each blank inside the sentence.
   - **Reading passages**: Wrap in a Markdown blockquote (`>` lines) and label as "Passage:".
   - **Short answer / proofreading / sentence transformation**: After the question, add 2-3 empty answer lines using ONLY underscores, one per line. NO BRACKETS, NO BRACES, NO PARENTHESES:
     ```
     ________________________________________________________________________________
     ________________________________________________________________________________
     ________________________________________________________________________________
     ```
   - **Composition / writing task**: Add 8-12 empty answer lines using ONLY underscores, 80 characters each. NO BRACKETS, NO BRACES, NO PARENTHESES, NO SQUARE BRACKETS:
     ```
     ________________________________________________________________________________
     ________________________________________________________________________________
     ________________________________________________________________________________
     ________________________________________________________________________________
     ________________________________________________________________________________
     ________________________________________________________________________________
     ________________________________________________________________________________
     ________________________________________________________________________________
     ```
   - **Cloze test**: Use `__________` (10 underscores) for each blank.
   - **Mathematics**: Use `__________` for answer blanks. For "Show your steps" questions, provide adequate space with underscore lines. Use **LaTeX math notation** for all mathematical expressions, wrapped with `$` for inline formulas and `$$` for display formulas. Examples: `$a = -3$`, `$(a - b)^2 \div (c + d)$`, `$$\frac{x^2 + y^2}{2}$$`, `$\sqrt{25}$`, `$x^2 + 3x - 4 = 0$`. This ensures proper rendering of mathematical symbols like exponents, fractions, square roots, and special characters.
   - **Images/Diagrams**: When a question requires a visual aid (such as diagrams, graphs, maps, or illustrations), use the special syntax `[IMAGE: descriptive prompt for the image]` instead of skipping the image. The prompt should be in English and describe exactly what the image should contain, including style, details, and any labels. Example: `[IMAGE: A simple bar chart showing monthly rainfall data for Hong Kong, with labels for months on X-axis and mm on Y-axis, clear grid lines, blue bars]`. Do NOT use standard Markdown image syntax `![alt](url)`.
   - **NO ASCII ART**: NEVER use text-based drawings, ASCII art, or character-based diagrams. ALWAYS use `[IMAGE: descriptive prompt]` instead. For example, instead of drawing a triangle with `/|\` characters, write: `[IMAGE: A right-angled triangle with base 6 cm, height 8 cm, and hypotenuse 10 cm, labeled with dimensions]`.
   - **No copyrighted material; write original passages.**

7. ANSWER KEY: every Part and every question must be answered. For multi-step math/science, show the steps. For writing tasks, give a model answer (1-3 sentences) or a clear marking scheme.

7.5. STUDENT WORKSPACE (CRITICAL - worksheets MUST have visible writing space, not just questions):
   - **Multiple choice**: no specific format required.
   - **Fill-in-the-blank (inline)**: Use `__________` (12+ underscores) for each blank inside the sentence.
   - **Short answer / sentence transformation**: Add 2-3 underscore lines after each question *if needed* (80 characters each, ONLY underscores).
   - **Reading comprehension short answer questions**: Add 2-3 underscore lines after each question (80 characters each, ONLY underscores).
   - **Composition / writing task (>=100 words required)**: Add 8-12 underscore lines (80 characters each, ONLY underscores).
   - **Mathematics show-your-steps**: Add 3-9 empty answer lines after the question.
   - **Cloze test**: Use `__________` (10 underscores) for each blank inline.
   - **CRITICAL RULE: USE ONLY UNDERSCORE CHARACTERS `_` FOR ANSWER LINES. DO NOT USE `[`, `]`, `{`, `}`, `(`, `)` OR ANY OTHER SYMBOLS.**


8. LANGUAGE:
   - If the teacher explicitly requests a Chinese version: **use 繁體中文 (Hong Kong convention) for ALL content including instructions, questions, and answer key.**
   - English subjects -> English (unless Chinese version requested)
   - Chinese subjects -> 繁體中文 (Hong Kong convention)
   - Other subjects -> English with bilingual key terms where appropriate (unless Chinese version requested)

Now produce the full worksheet. Begin with the title and follow the structure exactly. Do NOT include Design Rules or any other meta-content in the output."""

        # 替换占位符为实际值
        language_mode = "Chinese" if chinese_version else "English"
        system_prompt = (
            system_prompt
            .replace("__LEVEL__", level)
            .replace("__SUBJECT__", subject)
            .replace("__DIFFICULTY__", difficulty)
            .replace("__TIME__", str(suggested_time))
            .replace("__LANGUAGE_MODE__", language_mode)
            .replace("__OVERALL_NOTES__", overall_notes if overall_notes else "(none)")
        )

        # 附加课程文件上下文
        if curriculum_context:
            system_prompt += curriculum_context

        # 用户提示词
        user_prompt = (
            f"Generate a classwork worksheet for {level}, subject: {subject}, "
            f"difficulty: {difficulty}. Ensure it's appropriate for the Hong Kong education curriculum."
        )

        start_time = time.time()

        # 调用模型接口（非流式）
        response = get_model_client(model_name).chat.completions.create(
            model=get_model_name(model_name),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            max_tokens=get_model_max_tokens(model_name)
        )

        # 提取返回内容
        content = response.choices[0].message.content or ""

        # 过滤掉 <think>...</think> 思考块
        import re as _re
        content = _re.sub(r'<think>.*?</think>', '', content, flags=_re.DOTALL).strip()
        # 过滤 <thinking>...</thinking>
        content = _re.sub(r'<thinking>.*?</thinking>', '', content, flags=_re.DOTALL).strip()

        elapsed_time = time.time() - start_time

        print(f"[DEBUG] Generation time: {elapsed_time:.2f}s")
        print(f"[DEBUG] Content length: {len(content)} chars")

        # 提取 token 用量
        prompt_tokens = 0
        completion_tokens = 0
        total_tokens = 0
        if hasattr(response, 'usage') and response.usage:
            u = response.usage
            prompt_tokens = getattr(u, 'prompt_tokens', 0) or 0
            completion_tokens = getattr(u, 'completion_tokens', 0) or 0
            total_tokens = getattr(u, 'total_tokens', 0) or 0

        # 提取 finish_reason
        finish_reason = 'unknown'
        try:
            if hasattr(response, 'choices') and response.choices:
                first = response.choices[0]
                if hasattr(first, 'finish_reason') and first.finish_reason:
                    finish_reason = first.finish_reason
        except Exception:
            pass

        # 保存完成的任务
        _save_task(task_id, {
            "status": "completed",
            "content": content,
            "progress": 100,
            "level": level,
            "subject": subject,
            "elapsed": round(elapsed_time, 2),
            "tokens": {
                "prompt": prompt_tokens,
                "completion": completion_tokens,
                "total": total_tokens,
            },
            "finish_reason": finish_reason,
            "created_at": start_time,
        }, user)

        return jsonify({
            "content": content,
            "time": round(elapsed_time, 2),
            "tokens": {
                "prompt": prompt_tokens,
                "completion": completion_tokens,
                "total": total_tokens
            },
            "finish_reason": finish_reason,
            "task_id": task_id
        })

    except Exception as e:
        # 打印详细错误到控制台，便于调试
        print(f"[ERROR] 生成失败: {e}")
        return jsonify({"error": "生成失败，请稍后再试"}), 500


# ===========================================================================
# 自定义试卷生成接口（Custom Worksheet Builder）
# ===========================================================================
@requires_login
@app.route("/api/generate-custom", methods=["POST"])
def generate_custom_classwork():
    """
    用户自定义多个分部（Part）的试卷生成。
    请求体结构：
    {
      "level": "Primary 5",
      "subject": "English Language",
      "parts": [
        { "part_number": 1, "difficulty": "Basic",
          "question_type": "Multiple Choice", "weight": 2,
          "notes": "Focus on vocabulary" },
        ...
      ]
    }
    注意：weight 为相对数值（任意正数），后端按总和归一化到 100%。

    响应模式：
    - 默认（Accept: application/json）: 一次性返回完整 JSON（兼容旧前端）
    - Accept: text/event-stream: SSE 流式，每个 chunk 一个 data: 事件
    """
    data = request.get_json(silent=True) or {}
    level = (data.get("level") or "").strip()
    subject = (data.get("subject") or "").strip()
    parts = data.get("parts") or []
    curriculum_file = (data.get("curriculumFile") or "").strip()
    chinese_version = data.get("chineseVersion", False)
    overall_notes = (data.get("overallNotes") or "").strip()
    model_name = (data.get("model") or "").strip() or None

    # Load curriculum PDF context (if specified) so the AI aligns with
    # the Hong Kong curriculum when generating questions.
    curriculum_context = get_curriculum_context(curriculum_file)
    if curriculum_context:
        print(f"[INFO] [custom] Loaded curriculum: {curriculum_file}")

    # ---------- 校验 ----------
    if not level or not subject:
        return jsonify({"error": "Please select Level and Subject first."}), 400
    if not isinstance(parts, list) or len(parts) == 0:
        return jsonify({"error": "Please add at least one part."}), 400

    total_weight = 0.0
    for idx, p in enumerate(parts, 1):
        qtype = (p.get("question_type") or "").strip()
        weight = p.get("weight")
        difficulty = (p.get("difficulty") or "").strip()
        if not qtype:
            return jsonify({"error": f"Part {idx}: question type is required."}), 400
        if weight is None or str(weight).strip() == "":
            return jsonify({"error": f"Part {idx}: weight is required."}), 400
        try:
            weight_val = float(weight)
        except (TypeError, ValueError):
            return jsonify({"error": f"Part {idx}: weight must be a number."}), 400
        if weight_val <= 0:
            return jsonify({"error": f"Part {idx}: weight must be greater than 0."}), 400
        if not difficulty:
            return jsonify({"error": f"Part {idx}: difficulty is required."}), 400
        total_weight += weight_val

    if total_weight <= 0:
        return jsonify({"error": "Sum of weights must be greater than 0."}), 400

    # ---------- 归一化权重到 100% ----------
    normalized_parts = []
    parts_text_lines = []
    for p in parts:
        pn = p.get("part_number")
        diff = (p.get("difficulty") or "").strip()
        qtype = (p.get("question_type") or "").strip()
        weight_raw = float(p.get("weight"))
        notes = (p.get("notes") or "").strip() or "(none)"
        percent = round(weight_raw / total_weight * 100, 2)
        normalized_parts.append({
            "part_number": pn,
            "difficulty": diff,
            "question_type": qtype,
            "weight": percent,
            "notes": notes
        })
        parts_text_lines.append(
            f"- Part {pn}: difficulty = {diff}, question_type = {qtype}, "
            f"weight = {percent}% of total, notes = {notes}"
        )
    parts_text = "\n".join(parts_text_lines)

    language_mode = "Chinese (Hong Kong)" if chinese_version else "English"

    # 自定义试卷的 system prompt
    system_prompt = f"""You are an experienced Hong Kong education expert who specialises in designing worksheets that align with the Hong Kong local curriculum.

The teacher has requested a CUSTOM worksheet with the following configuration:

Overall defaults (used ONLY when no teacher note overrides them):
- Level: {level}
- Subject: {subject}
- Total marks: 100
- Time: 60 minutes (default)
- Marks per question: free to choose per part to fit the total
- Language: {language_mode}
- Overall Notes: {overall_notes if overall_notes else "(none)"}

Part breakdown (each part's marks = its weight % of 100):
{parts_text}
{curriculum_context}

==============================================================
PRIORITY HIERARCHY — READ THIS FIRST, IT OVERRIDES EVERYTHING ELSE
==============================================================
The following priority order applies to ALL instructions. Higher-priority instructions override lower-priority ones:

1. **Overall Notes** (highest priority): The teacher's overall notes apply to the ENTIRE worksheet and take precedence over everything else. If overall notes specify any requirements (e.g., topic focus, special instructions, format preferences, assessment criteria), ALL parts must adhere to those requirements.

2. **Part-specific Notes**: When a teacher's "notes" for a specific Part gives a specific instruction (number of questions, exact marks per question, time, scope, vocabulary, difficulty, format, etc.), that instruction **OVERRIDES** the system defaults and the CRITICAL RULES below. Specifically:
   - If notes say "5 questions" → write exactly 5, do NOT auto-expand to 8-15.
   - If notes say "2 marks each" → respect that, even if it doesn't match the weight %.
   - If notes say "30 minutes" or "90 minutes" → use that time, ignore the 60-minute default.
   - If notes specify a question count, do NOT add extra questions "to fill marks" — the total marks remain 100 across all parts, but per-question marks and question count follow the notes.
   - If notes specify scope (e.g. "only tenses", "only food vocabulary"), stay strictly within that scope.
   - If notes are EMPTY, follow the CRITICAL RULES below as best you can.

3. **System Defaults / CRITICAL RULES** (lowest priority): Apply only when neither Overall Notes nor Part-specific Notes provide instructions.

- DO NOT invent extra requirements, sections, decorations, or "extras" that the teacher did not ask for. Stay minimal and faithful.

==============================================================
CRITICAL RULES (apply when notes do not override them)
==============================================================
1. **Completeness**: Write ALL questions for each Part completely. DO NOT use "..." or "etc." to skip questions. Each question must be fully written out.
2. **Marks total per Part**: The sum of marks across questions in a Part should approximate the part's weight % of 100. Exact arithmetic is NOT required if notes give a different count; instead, follow the notes and label the part header with its weight-derived marks.
3. **Question quantity** (only when notes do NOT specify a count):
   - If a Part has 60-100% weight, write 8-15 questions.
   - If 30-59%, write 5-10 questions.
   - If 10-29%, write 3-6 questions.
   - If <10%, write 1-3 questions.
4. **Difficulty**: Match the difficulty specified for each Part (Basic / Advanced / Expert).
5. **Question types**: Strictly follow the "question_type" specified for each Part. If the type is "Multiple Choice" or includes MCQ, include A) B) C) D) options so students can circle the correct option with pen.

   **MCQ ANSWER-DISTRIBUTION RULES (MANDATORY — verify before outputting):**
   - Plan the answer-key letter sequence FIRST, before writing the questions, so you can balance it. Example for 8 MCQs in one Part: B, D, A, C, B, A, D, C (each option appears exactly twice; no 3-in-a-row; no obvious cycle).
   - For any Part with N MCQ questions, every option (A, B, C, D) must appear at least ⌈N/4⌉ times — no option may be missing.
   - No single option may appear in more than ~35% of the questions in a Part.
   - **No 3+ consecutive same letters** (e.g. AAA, BBB are FORBIDDEN).
   - **No simple repeating cycle** (e.g. ABCDABCD, ABABAB) — those are too predictable for students to guess.
   - After writing the questions AND the answer key, do a SELF-CHECK: count A/B/C/D in this Part. If any option is missing OR any option exceeds 35% OR there are 3+ in a row, SWAP the answer letter of one of the offending questions (and update its option text accordingly, e.g. change the option originally labeled "B" so it becomes the new correct one) until the distribution is balanced. Do this BEFORE finalising the output.
6. **Notes**: Strictly follow any subject-specific topics, vocabulary, or scope mentioned in the "notes" field of each Part. If notes specify "X questions" or "Y marks each", do not silently change those numbers.
7. **Answer lines**: Use ONLY underscore characters `_` for answer lines. Each answer line must be exactly 80 underscores long. NO square brackets, NO parentheses, NO braces.
8. **Reading passages**: When a Part requires a reading passage, default to 150-300 words for Primary, 250-400 words for Secondary. If the notes specify a length, follow the notes.
9. **Writing tasks**: If a Part is a writing task requiring ≥100 words, provide 8-12 answer lines (each 80 underscores) so students have enough space. If notes specify a target word count, use that.
10. **Math steps**: For calculation problems, provide clear step-by-step working in the answer key.

==============================================================
OUTPUT FORMAT (Markdown)
==============================================================
- Title: `# {subject} Custom Worksheet - {level}`
- For each Part: `## Part X: [type] (XX marks)` (XX = that part's weight-derived marks).
- After the title, include the questions in order.
- If a Part has a reading passage, use `### Passage` then the passage, then numbered questions.
- If a Part is MCQ, list A) B) C) D) and add `Answer: Ⓐ Ⓑ Ⓒ Ⓓ` after each question.
- All non-MCQ questions must be followed by 1-3 answer lines (80 underscores each) so students can write answers.
- **Mathematical Expressions**: Use **LaTeX math notation** for all mathematical expressions, wrapped with `$` for inline formulas and `$$` for display formulas. Examples: `$a = -3$`, `$(a - b)^2 \div (c + d)$`, `$$\frac{x^2 + y^2}{2}$$`, `$\sqrt{25}$`, `$x^2 + 3x - 4 = 0$`. This ensures proper rendering of mathematical symbols like exponents, fractions, square roots, and special characters.
- **Images/Diagrams**: When a question requires a visual aid (such as diagrams, graphs, maps, or illustrations), use the special syntax `[IMAGE: descriptive prompt for the image]` instead of skipping the image. The prompt should be in English and describe exactly what the image should contain, including style, details, and any labels. Example: `[IMAGE: A simple bar chart showing monthly rainfall data for Hong Kong, with labels for months on X-axis and mm on Y-axis, clear grid lines, blue bars]`. Do NOT use standard Markdown image syntax `![alt](url)`.
- **NO ASCII ART**: NEVER use text-based drawings, ASCII art, or character-based diagrams. ALWAYS use `[IMAGE: descriptive prompt]` instead. For example, instead of drawing a triangle with `/|\` characters, write: `[IMAGE: A right-angled triangle with base 6 cm, height 8 cm, and hypotenuse 10 cm, labeled with dimensions]`.
- At the very end, add `## Answers` and provide concise, complete answers for ALL questions in order.
- DO NOT include any sections after "## Answers".

ANSWER KEY REQUIREMENTS:
- For MCQ, show ONLY the letter (e.g. "Part 1, Q1: A").
- For calculation, show step-by-step working.
- For short answer, show the expected answer in 1-2 sentences.
- For writing tasks, show a model answer or key points.

==============================================================
FINAL REVIEW CHECKLIST — DO THIS AFTER WRITING EVERYTHING
==============================================================
Before you output the final response, mentally walk through this checklist and FIX any issues. Do not skip this step.

1. **Notes fidelity**: For each Part, re-read the teacher's notes. Does the question count match? Does the topic match? Are marks/time/scope as specified? If anything diverges, fix it.
2. **Completeness**: Count the questions in each Part. Is every question fully written (no "...", "etc.", or placeholder text)? Fix any truncations.
3. **MCQ distribution** (for any MCQ Part): Count A/B/C/D. Is every option present? Any option > 35%? Any 3+ in a row? Any obvious cycle? SWAP answer letters (and option text) to fix.
4. **Marks arithmetic**: Do the per-question marks in each Part add up to that Part's weight %? If not, adjust the per-question marks (NOT the question count, if notes specified a count).
5. **Answer key coverage**: For every question in every Part, is there a corresponding answer in `## Answers`? Fix any missing answers.
6. **Answer line format**: Every non-MCQ question has 1-3 underscore answer lines (80 underscores each, no brackets/braces)?
7. **No extra sections / no meta commentary**: The output contains only: title → parts → `## Answers`. No "Here's the worksheet:", no "I hope this helps", no closing remarks, no apologies, no advice.
8. **Hong Kong context**: Vocabulary, names, currency (HKD), date format, and cultural references should fit local Hong Kong primary/secondary classrooms.
9. **Language mode**: If the Language setting is "Chinese (Hong Kong)", verify ALL content (title, instructions, questions, answer key) is in 繁體中文. If "English", verify everything is in English.

Only AFTER all 9 items pass should you finalise the output.
"""
    user_prompt = f"Please generate the custom worksheet for {subject} at {level} with the {len(parts)} parts configured above. Honour the teacher's notes strictly. After writing, perform the FINAL REVIEW CHECKLIST."

    # ---------- 决定是否走流式 ----------
    wants_stream = (
        request.headers.get("Accept", "").startswith("text/event-stream")
        or (data.get("stream") is True)
    )

    if wants_stream:
        task_id = str(_uuid.uuid4())
        user = session.get("user", {}).get("username", "")
        print(f"[DEBUG] Custom stream request: level={level}, subject={subject}, user={user}, session_user={session.get('user')}")
        _save_task(task_id, {
            "status": "queued",
            "content": "",
            "progress": 0,
            "level": level,
            "subject": subject,
            "created_at": time.time(),
        }, user)
        return Response(
            _stream_custom_sse(system_prompt, user_prompt, task_id, model_name, user),
            mimetype="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",  # disable nginx buffering
                "Connection": "keep-alive",
            }
        )

    # ---------- 非流式路径（兼容旧前端） ----------
    return _non_stream_custom(system_prompt, user_prompt, normalized_parts, model_name)


def _stream_custom_sse(system_prompt: str, user_prompt: str, task_id: str = None, model_name: str = None, user: str = None):
    """SSE generator: yields 'data: {json}\\n\\n' for each chunk + final event.
    Buffers chunks to reduce SSE event frequency (avoids flooding the browser
    with hundreds of micro-render events that stall the main thread).

    Also emits a keep-alive ping every KEEPALIVE_INTERVAL seconds when no text
    has been flushed, so the browser can confirm the stream is still alive
    (and won't drop the connection) even when the model is slow.

    If task_id is provided, saves progress to disk for persistence.
    If model_name is provided, uses the specified model client.
    """
    import time
    import json as _json
    import re as _re

    start_time = time.time()
    print(f"[DEBUG] Custom SSE stream started (task={task_id})")

    full_content = ""
    prompt_tokens = 0
    completion_tokens = 0
    total_tokens = 0
    finish_reason = 'unknown'

    # ---- chunk buffering config ----
    BUFFER_SIZE = 192          # flush every N characters
    BUFFER_INTERVAL = 0.20     # OR every N seconds, whichever first
    KEEPALIVE_INTERVAL = 5.0   # ping every N seconds if no flush
    SAVE_INTERVAL = 2.0        # save progress every N seconds
    buffered = ""
    last_flush = time.time()
    last_save = time.time()

    def _flush_buffer():
        """Emit buffered text as one SSE chunk event (if non-empty)."""
        nonlocal buffered, last_flush
        if not buffered:
            return None
        payload = _json.dumps({"text": buffered, "char_count": len(full_content), "task_id": task_id})
        buffered = ""
        last_flush = time.time()
        return f"event: chunk\ndata: {payload}\n\n"

    def _save_progress():
        """Save current progress to disk (if task_id is set)."""
        nonlocal last_save
        if not task_id:
            return
        now = time.time()
        if now - last_save < SAVE_INTERVAL:
            return
        try:
            _save_task(task_id, {
                "status": "generating",
                "content": full_content,
                "progress": min(95, (len(full_content) / 14000) * 100),
                "elapsed": round(now - start_time, 2),
                "created_at": start_time,
            }, user)
            last_save = now
        except Exception:
            pass

    def _keepalive():
        """Emit a no-op ping so the browser knows the stream is alive."""
        return ": keep-alive\n\n"

    try:
        if task_id:
            _save_task(task_id, {
                "status": "started",
                "content": "",
                "progress": 0,
                "elapsed": 0,
                "created_at": start_time,
            }, user)
        yield f"event: start\ndata: {_json.dumps({'task_id': task_id})}\n\n"

        response = get_model_client(model_name).chat.completions.create(
            model=get_model_name(model_name),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            max_tokens=get_model_max_tokens(model_name),
            stream=True,
            stream_options={"include_usage": True} if "abab" not in get_model_name(model_name) else None,
        )

        for chunk in response:
            # Extract text delta
            try:
                choice = chunk.choices[0] if chunk.choices else None
                if choice and getattr(choice, 'delta', None) and getattr(choice.delta, 'content', None):
                    delta = choice.delta.content
                    full_content += delta
                    buffered += delta
                if choice and getattr(choice, 'finish_reason', None):
                    finish_reason = choice.finish_reason
            except Exception:
                pass

            # Extract usage
            try:
                if hasattr(chunk, 'usage') and chunk.usage:
                    prompt_tokens = getattr(chunk.usage, 'prompt_tokens', 0) or prompt_tokens
                    completion_tokens = getattr(chunk.usage, 'completion_tokens', 0) or completion_tokens
                    total_tokens = getattr(chunk.usage, 'total_tokens', 0) or total_tokens
            except Exception:
                pass

            now = time.time()
            if len(buffered) >= BUFFER_SIZE or (now - last_flush) >= BUFFER_INTERVAL:
                evt = _flush_buffer()
                if evt:
                    yield evt
                _save_progress()
            elif (now - last_flush) >= KEEPALIVE_INTERVAL:
                # No text flushed for a while → send keep-alive so the
                # browser's fetch doesn't drop the connection.
                yield _keepalive()
                _save_progress()

        # Flush any remaining buffered text
        evt = _flush_buffer()
        if evt:
            yield evt
        buffered = ""
        last_flush = time.time()

        # Cleanup <think> blocks
        clean_content = _re.sub(r'<think>.*?</think>', '', full_content, flags=_re.DOTALL).strip()
        clean_content = _re.sub(r'<thinking>.*?</thinking>', '', clean_content, flags=_re.DOTALL).strip()

        elapsed = time.time() - start_time
        print(f"[DEBUG] Custom SSE done in {elapsed:.2f}s, {len(clean_content)} chars, "
              f"{completion_tokens} completion tokens, finish={finish_reason}")

        done_payload = _json.dumps({
            "content": clean_content,
            "time": round(elapsed, 2),
            "tokens": {
                "prompt": prompt_tokens,
                "completion": completion_tokens,
                "total": total_tokens,
            },
            "finish_reason": finish_reason,
        })
        yield f"event: done\ndata: {done_payload}\n\n"

        # Save completed task
        if task_id:
            try:
                _save_task(task_id, {
                    "status": "completed",
                    "content": clean_content,
                    "progress": 100,
                    "elapsed": round(elapsed, 2),
                    "tokens": {
                        "prompt": prompt_tokens,
                        "completion": completion_tokens,
                        "total": total_tokens,
                    },
                    "finish_reason": finish_reason,
                    "created_at": start_time,
                }, user)
            except Exception:
                pass

    except Exception as e:
        print(f"[ERROR] Custom SSE stream failed: {e}")
        err_payload = _json.dumps({"error": "Generation failed. Please try again."})
        yield f"event: error\ndata: {err_payload}\n\n"

        # Save failed task
        if task_id:
            try:
                _save_task(task_id, {
                    "status": "failed",
                    "content": full_content,
                    "progress": min(95, (len(full_content) / 14000) * 100),
                    "elapsed": round(time.time() - start_time, 2),
                    "error": str(e),
                    "created_at": start_time,
                }, user)
            except Exception:
                pass


def _non_stream_custom(system_prompt, user_prompt, normalized_parts, model_name=None):
    """Original non-streaming path — kept for backward compatibility."""
    import time
    import re as _re
    try:
        start_time = time.time()
        response = get_model_client(model_name).chat.completions.create(
            model=get_model_name(model_name),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            max_tokens=get_model_max_tokens(model_name)
        )

        content = response.choices[0].message.content or ""
        content = _re.sub(r'<think>.*?</think>', '', content, flags=_re.DOTALL).strip()
        content = _re.sub(r'<thinking>.*?</thinking>', '', content, flags=_re.DOTALL).strip()

        elapsed = time.time() - start_time
        print(f"[DEBUG] Custom generation time: {elapsed:.2f}s, content length: {len(content)} chars")

        prompt_tokens = 0
        completion_tokens = 0
        total_tokens = 0
        if hasattr(response, 'usage') and response.usage:
            u = response.usage
            prompt_tokens = getattr(u, 'prompt_tokens', 0) or 0
            completion_tokens = getattr(u, 'completion_tokens', 0) or 0
            total_tokens = getattr(u, 'total_tokens', 0) or 0

        finish_reason = 'unknown'
        try:
            if hasattr(response, 'choices') and response.choices:
                first = response.choices[0]
                if hasattr(first, 'finish_reason') and first.finish_reason:
                    finish_reason = first.finish_reason
        except Exception:
            pass

        return jsonify({
            "content": content,
            "time": round(elapsed, 2),
            "tokens": {
                "prompt": prompt_tokens,
                "completion": completion_tokens,
                "total": total_tokens
            },
            "finish_reason": finish_reason,
            "weights_normalized": normalized_parts
        })

    except Exception as e:
        print(f"[ERROR] Custom generation failed: {e}")
        return jsonify({"error": "Generation failed. Please try again."}), 500


# ===========================================================================
# AI Polish Notes: refine a Part's notes via the LLM
# ===========================================================================
@requires_login
@app.route("/api/ai/polish-notes", methods=["POST"])
def polish_notes():
    """
    接收 { level, subject, question_type, notes, difficulty? }，
    让 AI 润色 notes（使其更具体、更专业、更适合该年级/科目/题型），
    返回润色后的 notes 文本。
    """
    try:
        data = request.get_json(silent=True) or {}
        level = (data.get("level") or "").strip()
        subject = (data.get("subject") or "").strip()
        qtype = (data.get("question_type") or "").strip()
        notes = (data.get("notes") or "").strip()
        difficulty = (data.get("difficulty") or "").strip() or "Advanced"
        model_name = (data.get("model") or "").strip() or None

        if not level or not subject or not qtype:
            return jsonify({"error": "Missing level / subject / question_type."}), 400
        if not notes:
            return jsonify({"error": "Notes cannot be empty."}), 400

        system_prompt = """You are an experienced Hong Kong education expert who helps teachers polish worksheet instructions.

Your job: take a teacher's rough notes for a worksheet part and rewrite them into a clear, specific, professional prompt instruction for an AI worksheet generator.

Rules:
1. Keep the teacher's original intent and topics.
2. Make the language clear, specific, and actionable.
3. Add concrete details that fit the Hong Kong curriculum for the given level and subject (vocabulary scope, grammar focus, question count if missing, difficulty, etc.).
4. Output ONLY the polished notes — no preamble, no quotes, no explanation.
5. Keep it concise: 1-3 sentences, max 200 words.
6. Use English."""

        user_prompt = f"""Level: {level}
Subject: {subject}
Question Type: {qtype}
Difficulty: {difficulty}
Original Notes: {notes}

Please polish the notes above."""

        import time
        start_time = time.time()
        response = get_model_client(model_name).chat.completions.create(
            model=get_model_name(model_name),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2,
            max_tokens=1024
        )

        polished = (response.choices[0].message.content or "").strip()
        # 过滤思考块
        import re as _re
        polished = _re.sub(r'<think>.*?</think>', '', polished, flags=_re.DOTALL).strip()
        polished = _re.sub(r'<thinking>.*?</thinking>', '', polished, flags=_re.DOTALL).strip()
        # 去掉首尾引号（AI 偶尔会包起来）
        polished = polished.strip('"').strip("'").strip()

        elapsed = time.time() - start_time
        print(f"[DEBUG] Polish notes time: {elapsed:.2f}s, length: {len(polished)} chars")

        return jsonify({
            "polished": polished,
            "time": round(elapsed, 2)
        })

    except Exception as e:
        print(f"[ERROR] Polish notes failed: {e}")
        return jsonify({"error": "Polishing failed. Please try again."}), 500




# DOCX 导出接口（延迟导入 docx，避免启动时的包冲突）
@requires_login
@app.route("/api/export/docx", methods=["POST"])
def export_docx():
    try:
        data = request.get_json(silent=True) or {}
        content = (data.get("content") or "").strip()
        level = (data.get("level") or "").strip()
        subject = (data.get("subject") or "").strip()

        if not content:
            return jsonify({"error": "缺少 content 参数"}), 400

        # 延迟导入：只在需要时才加载 docx，避免启动时的包冲突
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as import_err:
            if "exceptions" in str(import_err):
                return jsonify({
                    "error": (
                        "DOCX 导出失败：检测到旧版 `docx` 包冲突。\n"
                        "请手动执行以下命令修复：\n"
                        "  pip uninstall docx -y\n"
                        "  pip install python-docx"
                    )
                }), 500
            raise

        doc = Document()

        sections = doc.sections
        for section in sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5

        def strip_markdown(text):
            """去除文本中的 markdown 标记，保留纯文本。"""
            t = text
            # 临时保护 [IMAGE: ...] 标记，避免被其他正则表达式误处理
            image_markers = []
            def protect_image(match):
                image_markers.append(match.group(0))
                return f"__IMAGE_MARKER_{len(image_markers)-1}__"
            t = re.sub(r'\[IMAGE:[^\]]+\]', protect_image, t)
            
            # 代码块 ```
            t = re.sub(r'```[^`]*```', '', t, flags=re.DOTALL)
            t = t.replace('```', '')
            # 行内代码 `code`
            t = re.sub(r'`([^`]+)`', r'\1', t)
            # 加粗 **text** 或 __text__
            t = re.sub(r'\*\*([^*]+)\*\*', r'\1', t)
            t = re.sub(r'__([^_]+)__', r'\1', t)
            # 斜体 *text* 或 _text_
            t = re.sub(r'\*([^*\n]+)\*', r'\1', t)
            # 仅在形如"空白-_-文本-_-空白/标点"的边界识别 Markdown 斜体，
            # 避免误删 LaTeX 已转为 Unicode 后残留的下划线、或普通 a_b_c 等字面下划线。
            t = re.sub(r'(?:^|[\s(（\[\{])_([^_\n]+)_(?:$|[\s,，。\.！!；;：:？\)\]）\}）])', r'\1', t)
            # 链接 [text](url)
            t = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', t)
            # 图片 ![alt](url)
            t = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', t)
            # 标题标记 # （如果残留）
            t = re.sub(r'^#{1,6}\s*', '', t)
            # 引用标记 > （如果残留）
            t = re.sub(r'^>\s*', '', t)
            # 列表标记 - 或 * （行首）
            t = re.sub(r'^[\-\*]\s+', '', t)
            # 水平线 ---
            t = re.sub(r'^---+$', '', t)
            # 恢复 [IMAGE: ...] 标记
            for i, marker in enumerate(image_markers):
                t = t.replace(f"__IMAGE_MARKER_{i}__", marker)
            # 残留的多余空白
            t = t.strip()
            return t

        # Unicode 上标 / 下标映射表（Times New Roman、Calibri 等 docx 内置字体均支持）
        _SUPER_MAP = {
            '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
            '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
            '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾',
            'n': 'ⁿ', 'i': 'ⁱ', 'a': 'ᵃ', 'b': 'ᵇ', 'c': 'ᶜ',
            'd': 'ᵈ', 'e': 'ᵉ', 'f': 'ᶠ', 'g': 'ᵍ', 'h': 'ʰ',
            'j': 'ʲ', 'k': 'ᵏ', 'l': 'ˡ', 'm': 'ᵐ', 'o': 'ᵒ',
            'p': 'ᵖ', 'r': 'ʳ', 's': 'ˢ', 't': 'ᵗ', 'u': 'ᵘ',
            'v': 'ᵛ', 'w': 'ʷ', 'x': 'ˣ', 'y': 'ʸ', 'z': 'ᶻ',
        }
        _SUB_MAP = {
            '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
            '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
            '+': '₊', '-': '₋', '=': '₌', '(': '₍', ')': '₎',
            'a': 'ₐ', 'e': 'ₑ', 'h': 'ₕ', 'i': 'ᵢ', 'j': 'ⱼ',
            'k': 'ₖ', 'l': 'ₗ', 'm': 'ₘ', 'n': 'ₙ', 'o': 'ₒ',
            'p': 'ₚ', 'r': 'ᵣ', 's': 'ₛ', 't': 'ₜ', 'u': 'ᵤ',
            'v': 'ᵥ', 'x': 'ₓ',
        }

        def _to_super(s):
            return ''.join(_SUPER_MAP.get(ch, ch) for ch in s)

        def _to_sub(s):
            return ''.join(_SUB_MAP.get(ch, ch) for ch in s)

        def convert_latex(text):
            """将LaTeX数学公式转换为Unicode可读文本，用于docx导出。
            必须先于 strip_markdown 调用，避免下划线斜体规则误删 LaTeX 下标。
            """
            t = text

            # 1) 暂存 \$ 转义（LaTeX 中的美元符号）→ 占位符，处理完 LaTeX 后还原
            #    不能直接换成 $，否则下一步会把 $350 里的 $ 当成定界符删除
            t = t.replace(r'\$', '\x00DOLLAR\x00')
            # 2) 移除 $ 定界符（剩余的 $ 通常是行内/块公式定界符）
            t = t.replace('$$', '')
            t = t.replace('$', '')

            # 2) 不可见空格 / 间距命令
            t = re.sub(r'\\[,;:!]', ' ', t)          # \, \; \: \!
            t = re.sub(r'\\quad|\\qquad', '  ', t)    # \quad \qquad

            # 3) 复合 LaTeX 符号（先匹配长的，避免被短前缀吃掉）
            t = re.sub(r'\\times', '×', t)
            t = re.sub(r'\\div', '÷', t)
            t = re.sub(r'\\cdot', '·', t)
            t = re.sub(r'\\pm', '±', t)
            t = re.sub(r'\\mp', '∓', t)
            t = re.sub(r'\\leq', '≤', t)
            t = re.sub(r'\\le(?![a-zA-Z])', '≤', t)   # \le 但不匹配 \left 等
            t = re.sub(r'\\geq', '≥', t)
            t = re.sub(r'\\ge(?![a-zA-Z])', '≥', t)
            t = re.sub(r'\\neq', '≠', t)
            t = re.sub(r'\\ne(?![a-zA-Z])', '≠', t)
            t = re.sub(r'\\approx', '≈', t)
            t = re.sub(r'\\equiv', '≡', t)
            t = re.sub(r'\\propto', '∝', t)
            t = re.sub(r'\\rightarrow|\\to', '→', t)
            t = re.sub(r'\\Rightarrow', '⇒', t)
            t = re.sub(r'\\Leftarrow', '⇐', t)
            t = re.sub(r'\\Leftrightarrow', '⇔', t)
            t = re.sub(r'\\cup', '∪', t)
            t = re.sub(r'\\cap', '∩', t)
            t = re.sub(r'\\subset', '⊂', t)
            t = re.sub(r'\\subseteq', '⊆', t)
            t = re.sub(r'\\supset', '⊃', t)
            t = re.sub(r'\\supseteq', '⊇', t)
            t = re.sub(r'\\in(?![a-zA-Z])', '∈', t)   # 避免吃掉 \inf
            t = re.sub(r'\\notin', '∉', t)
            t = re.sub(r'\\sum', 'Σ', t)
            t = re.sub(r'\\prod', 'Π', t)
            t = re.sub(r'\\int', '∫', t)
            t = re.sub(r'\\partial', '∂', t)
            t = re.sub(r'\\nabla', '∇', t)
            t = re.sub(r'\\infty', '∞', t)
            t = re.sub(r'\\angle', '∠', t)
            t = re.sub(r'\\degree|^{\\circ}', '°', t)

            # 4) 希腊字母（小写）
            for src, dst in [
                ('\\alpha', 'α'), ('\\beta', 'β'), ('\\gamma', 'γ'),
                ('\\delta', 'δ'), ('\\epsilon', 'ε'), ('\\zeta', 'ζ'),
                ('\\eta', 'η'), ('\\theta', 'θ'), ('\\iota', 'ι'),
                ('\\kappa', 'κ'), ('\\lambda', 'λ'), ('\\mu', 'μ'),
                ('\\nu', 'ν'), ('\\xi', 'ξ'), ('\\omicron', 'ο'),
                ('\\pi', 'π'), ('\\rho', 'ρ'), ('\\sigma', 'σ'),
                ('\\tau', 'τ'), ('\\upsilon', 'υ'), ('\\phi', 'φ'),
                ('\\varphi', 'ϕ'), ('\\chi', 'χ'), ('\\psi', 'ψ'),
                ('\\omega', 'ω'),
            ]:
                t = t.replace(src, dst)
            # 希腊字母（大写）
            for src, dst in [
                ('\\Gamma', 'Γ'), ('\\Delta', 'Δ'), ('\\Theta', 'Θ'),
                ('\\Lambda', 'Λ'), ('\\Xi', 'Ξ'), ('\\Pi', 'Π'),
                ('\\Sigma', 'Σ'), ('\\Phi', 'Φ'), ('\\Psi', 'Ψ'),
                ('\\Omega', 'Ω'),
            ]:
                t = t.replace(src, dst)

            # 5) 分数 \frac{a}{b} -> a/b（必须用 regex，不能先 replace \frac 为空格）
            t = re.sub(r'\\frac\s*\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}\s*\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}', r'\1/\2', t)
            t = re.sub(r'\\dfrac\s*\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}\s*\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}', r'\1/\2', t)

            # 6) 平方根 \sqrt{x} 或 \sqrt[n]{x} -> √(x) 或 ⁿ√(x)
            t = re.sub(r'\\sqrt\[([^\]]+)\]\s*\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}', lambda m: _to_super(m.group(1)) + '√(' + m.group(2) + ')', t)
            t = re.sub(r'\\sqrt\s*\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}', r'√(\1)', t)

            # 7) 上标 ^{n} / ^n -> Unicode 上标（在删大括号之前）
            t = re.sub(r'\^\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}', lambda m: _to_super(m.group(1)), t)
            t = re.sub(r'\^([0-9n+\-=()a-zA-Z])', lambda m: _to_super(m.group(1)), t)

            # 8) 下标 _{n} / _n -> Unicode 下标（必须在 strip_markdown 之前执行！）
            t = re.sub(r'_\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}', lambda m: _to_sub(m.group(1)), t)
            t = re.sub(r'_([0-9+\-=()a-zA-Z])', lambda m: _to_sub(m.group(1)), t)

            # 9) 清理残留的 { }
            t = t.replace('{', '').replace('}', '')

            # 10) 还原 \$ 占位符为美元符号
            t = t.replace('\x00DOLLAR\x00', '$')

            # 11) 折叠多余空白
            t = re.sub(r' {2,}', ' ', t)
            return t

        def process_text(text):
            """统一处理文本：先转换LaTeX公式（把下划线/上标变为Unicode），
            再去除markdown标记，避免strip_markdown的下划线斜体规则误删LaTeX下标。
            """
            return strip_markdown(convert_latex(text))

        lines = content.split("\n")

        for line in lines:
            stripped = line.strip()

            if stripped.startswith("# "):
                p = doc.add_paragraph()
                run = p.add_run(process_text(stripped[2:]))
                run.font.size = Pt(16)
                run.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)

            elif stripped.startswith("## "):
                p = doc.add_paragraph()
                run = p.add_run(process_text(stripped[3:]))
                run.font.size = Pt(14)
                run.font.bold = True
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(8)

            elif stripped.startswith("### "):
                p = doc.add_paragraph()
                run = p.add_run(process_text(stripped[4:]))
                run.font.size = Pt(12)
                run.font.bold = True
                p.paragraph_format.space_before = Pt(8)

            elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
                p = doc.add_paragraph()
                run = p.add_run(process_text(stripped[2:-2]))
                run.font.bold = True

            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(process_text(stripped[2:]), style="List Bullet")
                p.paragraph_format.space_after = Pt(2)

            elif re.match(r'^(?:Q(?:uestion)?\s*)?\d+\s*[\.\)\u3001]\s*', stripped, re.IGNORECASE):
                body = re.sub(r'^(?:Q(?:uestion)?\s*)?\d+\s*[\.\)\u3001]\s*', '', stripped, count=1, flags=re.IGNORECASE).strip()
                if body:
                    p = doc.add_paragraph(process_text(body), style="List Number")
                    p.paragraph_format.space_after = Pt(2)
                else:
                    p = doc.add_paragraph(process_text(stripped))

            elif stripped.startswith("[____]"):
                p = doc.add_paragraph()
                run = p.add_run("_" * 80)
                run.font.size = Pt(12)
                p.paragraph_format.space_after = Pt(4)

            elif stripped.startswith("_") and len(stripped) >= 60 and all(c == "_" for c in stripped):
                p = doc.add_paragraph()
                run = p.add_run("_" * 80)
                run.font.size = Pt(12)
                p.paragraph_format.space_after = Pt(4)

            elif stripped.startswith("> "):
                p = doc.add_paragraph()
                run = p.add_run(process_text(stripped[2:]))
                p.paragraph_format.left_indent = Cm(1)

            elif "Answer:" in stripped and "Ⓐ" in stripped and "Ⓓ" in stripped:
                p = doc.add_paragraph()
                run = p.add_run(process_text(stripped))
                run.font.size = Pt(12)
                p.paragraph_format.space_after = Pt(8)

            elif stripped.startswith("---"):
                p = doc.add_paragraph()
                run = p.add_run("-" * 60)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif stripped == "":
                doc.add_paragraph()

            else:
                p = doc.add_paragraph(process_text(stripped))
                p.paragraph_format.space_after = Pt(4)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        safe_level = level.replace(" ", "_") if level else "worksheet"
        safe_subject = subject.replace(" ", "_") if subject else "class"
        filename = f"Classwork_{safe_level}_{safe_subject}.docx"

        return send_file(
            buffer,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print(f"[ERROR] DOCX 导出失败: {e}")
        return jsonify({"error": "DOCX 导出失败，请稍后再试"}), 500


# 测试 MiniMax M3 模型调用
@app.route("/api/test/minimax", methods=["GET", "POST"])
def test_minimax():
    """测试 MiniMax M3 模型调用是否成功"""
    try:
        # 优先使用请求中的模型名，否则用 .env 里的，最后用默认
        if request.method == "POST":
            data = request.get_json(silent=True) or {}
            test_model = data.get("model") or os.getenv("MODEL_NAME", "MiniMax-Text-01")
        else:
            test_model = os.getenv("MODEL_NAME", "MiniMax-Text-01")

        api_key = os.getenv("MINIMAX_API_KEY")
        if not api_key:
            return jsonify({
                "success": False,
                "error": "未找到 MINIMAX_API_KEY 环境变量"
            }), 400

        # 初始化客户端（复用同一个 client）
        client = OpenAI(
            api_key=api_key,
            base_url="https://api.minimax.chat/v1"
        )

        # 简单测试 prompt
        test_prompt = "Reply with: Hello from MiniMax. Model is working."

        print(f"[TEST] Testing model: {test_model}")
        start = time.time()
        response = client.chat.completions.create(
            model=test_model,
            messages=[
                {"role": "user", "content": test_prompt}
            ],
            temperature=0.3,
            max_tokens=50
        )
        elapsed = time.time() - start

        # 提取响应
        result = {
            "success": True,
            "model": test_model,
            "elapsed": round(elapsed, 2),
            "raw_response_type": str(type(response).__name__),
            "raw_response": str(response)[:500]
        }

        # 尝试提取内容
        try:
            if hasattr(response, 'choices') and response.choices:
                first = response.choices[0]
                if hasattr(first, 'message'):
                    result["content"] = first.message.content
                # 提取 finish_reason
                if hasattr(first, 'finish_reason'):
                    result["finish_reason"] = first.finish_reason
        except Exception as e:
            result["extract_error"] = str(e)

        # 尝试提取 usage
        try:
            if hasattr(response, 'usage') and response.usage:
                u = response.usage
                result["usage"] = {
                    "prompt_tokens": getattr(u, 'prompt_tokens', None),
                    "completion_tokens": getattr(u, 'completion_tokens', None),
                    "total_tokens": getattr(u, 'total_tokens', None)
                }
            elif hasattr(response, '_raw_response'):
                raw = response._raw_response.json()
                if 'usage' in raw:  
                    result["usage_raw"] = raw['usage']
        except Exception as e:
            result["usage_error"] = str(e)

        print(f"[TEST] Success: {result.get('content', '(no content)')[:100]}")
        return jsonify(result)

    except Exception as e:
        print(f"[TEST] Failed: {e}")
        return jsonify({
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }), 500


if __name__ == "__main__":
    _start_cleanup_thread()
    print("Classwork Studio server running on http://localhost:3000")
    app.run(host="0.0.0.0", port=3000, debug=False)
